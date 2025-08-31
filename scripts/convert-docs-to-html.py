#!/usr/bin/env python3
"""
Script para convertir documentos DOCX a HTML y extraer toda la información de cada caso
"""

import os
import json
import re
from pathlib import Path
from typing import Dict, List, Any
import subprocess
import sys

class DocxToHtmlConverter:
    def __init__(self):
        self.base_path = Path(__file__).parent.parent
        self.cases_path = self.base_path / "public" / "casos"
        self.output_path = self.base_path / "scripts" / "extracted_cases"
        
        # Crear directorio de salida si no existe
        self.output_path.mkdir(exist_ok=True)
        
    def install_dependencies(self):
        """Instala las dependencias necesarias"""
        print("📦 Instalando dependencias...")
        try:
            subprocess.run([sys.executable, "-m", "pip", "install", "python-docx", "beautifulsoup4"], check=True)
            print("✅ Dependencias instaladas correctamente")
        except subprocess.CalledProcessError as e:
            print(f"❌ Error instalando dependencias: {e}")
            return False
        return True
    
    def convert_docx_to_html(self, docx_path: Path) -> str:
        """Convierte un archivo DOCX a HTML"""
        try:
            from docx import Document
            from bs4 import BeautifulSoup
            
            # Leer el documento DOCX
            doc = Document(docx_path)
            
            # Crear HTML básico
            html_content = []
            html_content.append("<!DOCTYPE html>")
            html_content.append("<html lang='es'>")
            html_content.append("<head>")
            html_content.append("<meta charset='UTF-8'>")
            html_content.append(f"<title>{docx_path.stem}</title>")
            html_content.append("<style>")
            html_content.append("body { font-family: Arial, sans-serif; margin: 40px; line-height: 1.6; }")
            html_content.append("h1, h2, h3 { color: #333; }")
            html_content.append("p { margin-bottom: 10px; }")
            html_content.append(".section { margin-bottom: 30px; border-bottom: 1px solid #eee; padding-bottom: 20px; }")
            html_content.append("</style>")
            html_content.append("</head>")
            html_content.append("<body>")
            
            # Procesar cada párrafo
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if text:
                    # Detectar títulos por formato o contenido
                    if paragraph.style.name.startswith('Heading') or self.is_title(text):
                        html_content.append(f"<h2>{text}</h2>")
                    else:
                        html_content.append(f"<p>{text}</p>")
            
            # Procesar tablas si las hay
            for table in doc.tables:
                html_content.append("<table border='1' style='border-collapse: collapse; width: 100%; margin: 20px 0;'>")
                for row in table.rows:
                    html_content.append("<tr>")
                    for cell in row.cells:
                        html_content.append(f"<td style='padding: 8px; border: 1px solid #ddd;'>{cell.text}</td>")
                    html_content.append("</tr>")
                html_content.append("</table>")
            
            html_content.append("</body>")
            html_content.append("</html>")
            
            return "\n".join(html_content)
            
        except ImportError:
            print("❌ Error: No se pudo importar python-docx. Instalando...")
            if self.install_dependencies():
                return self.convert_docx_to_html(docx_path)
            else:
                return f"<html><body><h1>Error: No se pudo convertir {docx_path.name}</h1></body></html>"
        except Exception as e:
            print(f"❌ Error convirtiendo {docx_path.name}: {e}")
            return f"<html><body><h1>Error convirtiendo {docx_path.name}</h1><p>{str(e)}</p></body></html>"
    
    def is_title(self, text: str) -> bool:
        """Detecta si un texto es un título basándose en su contenido"""
        title_patterns = [
            r'^CASO\s+\d+',
            r'^Ficha\s+rápida',
            r'^Resumen$',
            r'^Hechos$',
            r'^Estrategia\s+de\s+defensa$',
            r'^Pruebas\s+clave$',
            r'^Resolución$',
            r'^Marco\s+legal$',
            r'^Imágenes\s*\(anexos\s+del\s+expediente\)$',
            r'^Enlaces\s+y\s+notas\s*\(citas\)$',
            r'^SEO\s+y\s+anclas\s*\(web\)$',
            r'^Nota\s+de\s+verificación$'
        ]
        
        for pattern in title_patterns:
            if re.match(pattern, text, re.IGNORECASE):
                return True
        return False
    
    def extract_case_info(self, html_content: str) -> Dict[str, Any]:
        """Extrae información estructurada del HTML"""
        from bs4 import BeautifulSoup
        
        soup = BeautifulSoup(html_content, 'html.parser')
        
        case_info = {
            'raw_html': html_content,
            'sections': {},
            'metadata': {}
        }
        
        # Extraer secciones
        current_section = None
        current_content = []
        
        for element in soup.find_all(['h2', 'p']):
            if element.name == 'h2':
                # Guardar sección anterior
                if current_section and current_content:
                    case_info['sections'][current_section] = '\n'.join(current_content).strip()
                
                # Nueva sección
                current_section = element.get_text().strip()
                current_content = []
            elif element.name == 'p' and current_section:
                current_content.append(element.get_text().strip())
        
        # Guardar última sección
        if current_section and current_content:
            case_info['sections'][current_section] = '\n'.join(current_content).strip()
        
        return case_info
    
    def process_all_cases(self):
        """Procesa todos los casos DOCX"""
        print("🔄 Procesando todos los casos DOCX...")
        
        # Buscar todos los archivos DOCX
        docx_files = list(self.cases_path.glob("*.docx"))
        
        if not docx_files:
            print("❌ No se encontraron archivos DOCX")
            return
        
        print(f"📁 Encontrados {len(docx_files)} archivos DOCX")
        
        all_cases = {}
        
        for docx_file in docx_files:
            print(f"📝 Procesando: {docx_file.name}")
            
            # Convertir a HTML
            html_content = self.convert_docx_to_html(docx_file)
            
            # Extraer información
            case_info = self.extract_case_info(html_content)
            
            # Guardar HTML
            html_file = self.output_path / f"{docx_file.stem}.html"
            with open(html_file, 'w', encoding='utf-8') as f:
                f.write(html_content)
            
            # Guardar información estructurada
            json_file = self.output_path / f"{docx_file.stem}.json"
            with open(json_file, 'w', encoding='utf-8') as f:
                json.dump(case_info, f, ensure_ascii=False, indent=2)
            
            all_cases[docx_file.stem] = case_info
            
            print(f"✅ {docx_file.name} procesado -> {html_file.name} y {json_file.name}")
        
        # Guardar resumen de todos los casos
        summary_file = self.output_path / "all_cases_summary.json"
        with open(summary_file, 'w', encoding='utf-8') as f:
            json.dump(all_cases, f, ensure_ascii=False, indent=2)
        
        print(f"🎉 Procesamiento completado!")
        print(f"📂 Archivos generados en: {self.output_path}")
        print(f"📊 Resumen guardado en: {summary_file}")
        
        return all_cases
    
    def create_index_html(self, cases: Dict[str, Any]):
        """Crea un archivo index.html con enlaces a todos los casos"""
        html_content = []
        html_content.append("<!DOCTYPE html>")
        html_content.append("<html lang='es'>")
        html_content.append("<head>")
        html_content.append("<meta charset='UTF-8'>")
        html_content.append("<title>Casos Extraídos - STANS ABOGADOS</title>")
        html_content.append("<style>")
        html_content.append("body { font-family: Arial, sans-serif; margin: 40px; line-height: 1.6; }")
        html_content.append("h1 { color: #333; text-align: center; }")
        html_content.append(".case-link { display: block; margin: 10px 0; padding: 15px; background: #f5f5f5; border-radius: 5px; text-decoration: none; color: #333; }")
        html_content.append(".case-link:hover { background: #e0e0e0; }")
        html_content.append(".case-number { font-weight: bold; color: #666; }")
        html_content.append("</style>")
        html_content.append("</head>")
        html_content.append("<body>")
        html_content.append("<h1>📋 Casos Extraídos - STANS ABOGADOS</h1>")
        
        # Ordenar casos por número
        sorted_cases = sorted(cases.keys(), key=lambda x: self.extract_case_number(x))
        
        for case_name in sorted_cases:
            case_number = self.extract_case_number(case_name)
            display_name = case_name.replace('_', ' ').replace('CASO', 'CASO ').replace('CASO  ', 'CASO ')
            
            html_content.append(f'<a href="{case_name}.html" class="case-link">')
            html_content.append(f'<span class="case-number">Caso {case_number}</span><br>')
            html_content.append(f'{display_name}')
            html_content.append('</a>')
        
        html_content.append("</body>")
        html_content.append("</html>")
        
        index_file = self.output_path / "index.html"
        with open(index_file, 'w', encoding='utf-8') as f:
            f.write('\n'.join(html_content))
        
        print(f"📄 Índice creado: {index_file}")
    
    def extract_case_number(self, case_name: str) -> int:
        """Extrae el número del caso del nombre del archivo"""
        match = re.search(r'CASO[_\s]*(\d+)', case_name, re.IGNORECASE)
        if match:
            return int(match.group(1))
        return 999  # Para casos sin número claro

def main():
    """Función principal"""
    converter = DocxToHtmlConverter()
    
    try:
        # Procesar todos los casos
        cases = converter.process_all_cases()
        
        if cases:
            # Crear índice HTML
            converter.create_index_html(cases)
            
            print("\n🎯 Resumen:")
            print(f"✅ {len(cases)} casos procesados")
            print(f"📁 Archivos HTML generados en: {converter.output_path}")
            print(f"🔗 Abre {converter.output_path}/index.html para ver todos los casos")
            
    except Exception as e:
        print(f"❌ Error durante el procesamiento: {e}")

if __name__ == "__main__":
    main()
