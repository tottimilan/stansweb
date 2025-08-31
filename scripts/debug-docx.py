import os
from docx import Document

# Listar todos los archivos DOCX
docx_files = [f for f in os.listdir('public/casos') if f.endswith('.docx')]
print(f"📁 Encontrados {len(docx_files)} archivos DOCX\n")

# Analizar el primer archivo para ver su estructura
if docx_files:
    test_file = docx_files[0]
    print(f"🔍 Analizando estructura de: {test_file}\n")
    
    doc = Document(os.path.join('public/casos', test_file))
    
    print("📄 Primeros 20 párrafos del documento:")
    print("=" * 70)
    
    for i, para in enumerate(doc.paragraphs[:20]):
        text = para.text.strip()
        if text:
            print(f"{i+1:3d}: {text[:100]}...")
    
    print("\n" + "=" * 70)
    print("\n📋 Buscando secciones principales en todos los archivos:\n")
    
    for filename in docx_files[:5]:  # Primeros 5 archivos
        print(f"\n📄 {filename}:")
        doc = Document(os.path.join('public/casos', filename))
        sections_found = []
        
        for para in doc.paragraphs[:50]:  # Primeros 50 párrafos
            text = para.text.strip()
            if text in ['FICHA RÁPIDA', 'RESUMEN', 'HECHOS', 'ESTRATEGIA DE DEFENSA', 
                       'PRUEBAS CLAVE', 'RESOLUCIÓN', 'MARCO LEGAL', 'IMÁGENES', 
                       'ENLACES Y NOTAS', 'SEO Y ANCLAS']:
                sections_found.append(text)
        
        print(f"   Secciones encontradas: {sections_found}")
